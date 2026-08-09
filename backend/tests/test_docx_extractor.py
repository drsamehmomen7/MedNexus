from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

from backend.app.modules.medical_document_intelligence.contracts.document_content import (
    DocumentContent,
)
from backend.app.modules.medical_document_intelligence.extractors.docx import (
    DocxDocumentExtractor,
)


def create_docx(
    path: Path,
    *,
    paragraphs: tuple[str, ...] = (),
    table_rows: tuple[tuple[str, ...], ...] = (),
    header_text: str | None = None,
    footer_text: str | None = None,
) -> Path:
    """
    Create a DOCX fixture for extractor tests.
    """

    document = Document()

    for paragraph_text in paragraphs:
        document.add_paragraph(paragraph_text)

    if table_rows:
        column_count = max(
            len(row)
            for row in table_rows
        )

        table = document.add_table(
            rows=len(table_rows),
            cols=column_count,
        )

        for row_index, row_values in enumerate(table_rows):
            for column_index, value in enumerate(row_values):
                table.cell(
                    row_index,
                    column_index,
                ).text = value

    if header_text is not None:
        document.sections[0].header.paragraphs[0].text = (
            header_text
        )

    if footer_text is not None:
        document.sections[0].footer.paragraphs[0].text = (
            footer_text
        )

    document.save(path)

    return path


def test_supported_extension():
    extractor = DocxDocumentExtractor()

    assert extractor.supported_extensions == (
        ".docx",
    )


def test_media_type():
    extractor = DocxDocumentExtractor()

    assert extractor.media_type == (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    )


def test_supports_docx_extension():
    extractor = DocxDocumentExtractor()

    assert extractor.supports(
        "report.docx"
    ) is True


def test_supports_uppercase_docx_extension():
    extractor = DocxDocumentExtractor()

    assert extractor.supports(
        "REPORT.DOCX"
    ) is True


def test_rejects_unsupported_extension():
    extractor = DocxDocumentExtractor()

    assert extractor.supports(
        "report.pdf"
    ) is False


def test_extract_returns_document_content(tmp_path):
    file_path = create_docx(
        tmp_path / "report.docx",
        paragraphs=(
            "Pathology Department",
            "Patient Name: Ahmed Hassan",
        ),
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    assert isinstance(
        result,
        DocumentContent,
    )


def test_extracts_body_paragraphs(tmp_path):
    file_path = create_docx(
        tmp_path / "report.docx",
        paragraphs=(
            "Pathology Department",
            "Patient Name: Ahmed Hassan",
            "Final Diagnosis",
        ),
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    assert "Pathology Department" in result.text
    assert "Patient Name: Ahmed Hassan" in result.text
    assert "Final Diagnosis" in result.text


def test_preserves_paragraph_order(tmp_path):
    file_path = create_docx(
        tmp_path / "report.docx",
        paragraphs=(
            "First paragraph",
            "Second paragraph",
            "Third paragraph",
        ),
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    first_position = result.text.index(
        "First paragraph"
    )
    second_position = result.text.index(
        "Second paragraph"
    )
    third_position = result.text.index(
        "Third paragraph"
    )

    assert first_position < second_position
    assert second_position < third_position


def test_extracts_table_content(tmp_path):
    file_path = create_docx(
        tmp_path / "laboratory.docx",
        table_rows=(
            (
                "Test",
                "Result",
                "Unit",
            ),
            (
                "Hemoglobin",
                "13.9",
                "g/dL",
            ),
            (
                "WBC",
                "11.8",
                "10^9/L",
            ),
        ),
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    assert "Test\tResult\tUnit" in result.text
    assert "Hemoglobin\t13.9\tg/dL" in result.text
    assert "WBC\t11.8\t10^9/L" in result.text


def test_preserves_body_paragraph_and_table_order(tmp_path):
    file_path = tmp_path / "ordered.docx"

    document = Document()

    document.add_paragraph(
        "Before table"
    )

    table = document.add_table(
        rows=1,
        cols=2,
    )

    table.cell(0, 0).text = "Test"
    table.cell(0, 1).text = "Result"

    document.add_paragraph(
        "After table"
    )

    document.save(
        file_path
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    before_position = result.text.index(
        "Before table"
    )
    table_position = result.text.index(
        "Test\tResult"
    )
    after_position = result.text.index(
        "After table"
    )

    assert before_position < table_position
    assert table_position < after_position


def test_extracts_header_text(tmp_path):
    file_path = create_docx(
        tmp_path / "report.docx",
        paragraphs=(
            "Body content",
        ),
        header_text="Ministry of Health",
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    assert "Ministry of Health" in result.text
    assert result.metadata[
        "header_block_count"
    ] == 1


def test_extracts_footer_text(tmp_path):
    file_path = create_docx(
        tmp_path / "report.docx",
        paragraphs=(
            "Body content",
        ),
        footer_text="Confidential Medical Report",
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    assert "Confidential Medical Report" in result.text
    assert result.metadata[
        "footer_block_count"
    ] == 1


def test_headers_and_footers_are_appended_after_body(tmp_path):
    file_path = create_docx(
        tmp_path / "report.docx",
        paragraphs=(
            "Body content",
        ),
        header_text="Header content",
        footer_text="Footer content",
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    body_position = result.text.index(
        "Body content"
    )
    header_position = result.text.index(
        "Header content"
    )
    footer_position = result.text.index(
        "Footer content"
    )

    assert body_position < header_position
    assert header_position < footer_position


def test_extracts_header_table(tmp_path):
    file_path = tmp_path / "header-table.docx"

    document = Document()

    header = document.sections[0].header

    table = header.add_table(
        rows=1,
        cols=2,
        width=0,
    )

    table.cell(0, 0).text = "Hospital"
    table.cell(0, 1).text = "Case Number"

    document.add_paragraph(
        "Body content"
    )

    document.save(
        file_path
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    assert "Hospital\tCase Number" in result.text


def test_extracts_footer_table(tmp_path):
    file_path = tmp_path / "footer-table.docx"

    document = Document()

    footer = document.sections[0].footer

    table = footer.add_table(
        rows=1,
        cols=2,
        width=0,
    )

    table.cell(0, 0).text = "Approved"
    table.cell(0, 1).text = "Verified"

    document.add_paragraph(
        "Body content"
    )

    document.save(
        file_path
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    assert "Approved\tVerified" in result.text


def test_extracts_source_metadata(tmp_path):
    file_path = create_docx(
        tmp_path / "clinical-report.docx",
        paragraphs=(
            "Clinical Report",
        ),
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    assert result.source_name == (
        "clinical-report.docx"
    )
    assert result.extension == ".docx"
    assert result.media_type == (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    )
    assert result.file_size > 0
    assert result.encoding is None
    assert result.page_count is None


def test_extracts_document_metadata_counts(tmp_path):
    file_path = create_docx(
        tmp_path / "report.docx",
        paragraphs=(
            "Paragraph one",
            "Paragraph two",
        ),
        table_rows=(
            (
                "Column 1",
                "Column 2",
            ),
        ),
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    assert result.metadata[
        "extractor"
    ] == "docx"

    assert result.metadata[
        "paragraph_count"
    ] == 2

    assert result.metadata[
        "table_count"
    ] == 1

    assert result.metadata[
        "section_count"
    ] == 1

    assert result.metadata[
        "contains_tables"
    ] is True

    assert result.metadata[
        "preserves_body_order"
    ] is True


def test_document_without_tables_reports_false(tmp_path):
    file_path = create_docx(
        tmp_path / "report.docx",
        paragraphs=(
            "Only text",
        ),
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    assert result.metadata[
        "contains_tables"
    ] is False

    assert result.metadata[
        "table_count"
    ] == 0


def test_empty_document_returns_warning(tmp_path):
    file_path = create_docx(
        tmp_path / "empty.docx",
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    assert result.is_empty is True
    assert result.has_warnings is True
    assert (
        "The DOCX document contains no extractable text."
        in result.warnings
    )


def test_empty_paragraphs_are_ignored(tmp_path):
    file_path = create_docx(
        tmp_path / "report.docx",
        paragraphs=(
            "",
            "   ",
            "Clinical content",
            "",
        ),
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    assert result.text == "Clinical content"


def test_normalizes_line_breaks_inside_paragraph(tmp_path):
    file_path = create_docx(
        tmp_path / "report.docx",
        paragraphs=(
            "Line one\nLine two",
        ),
    )

    extractor = DocxDocumentExtractor()

    result = extractor.extract(
        file_path
    )

    assert "Line one\nLine two" in result.text


def test_missing_file_raises_file_not_found():
    extractor = DocxDocumentExtractor()

    with pytest.raises(
        FileNotFoundError,
    ):
        extractor.extract(
            "missing-report.docx"
        )


def test_directory_path_is_rejected(tmp_path):
    extractor = DocxDocumentExtractor()

    with pytest.raises(
        ValueError,
        match="is not a file",
    ):
        extractor.extract(
            tmp_path
        )


def test_empty_string_path_is_rejected():
    extractor = DocxDocumentExtractor()

    with pytest.raises(
        ValueError,
        match="path must not be empty",
    ):
        extractor.extract(
            "   "
        )


def test_non_path_value_is_rejected():
    extractor = DocxDocumentExtractor()

    with pytest.raises(
        TypeError,
        match=(
            "path must be a string or pathlib.Path"
        ),
    ):
        extractor.extract(
            123
        )


def test_wrong_extension_is_rejected(tmp_path):
    file_path = tmp_path / "report.txt"

    file_path.write_text(
        "Not a DOCX document",
        encoding="utf-8",
    )

    extractor = DocxDocumentExtractor()

    with pytest.raises(
        ValueError,
        match="supports only '.docx' files",
    ):
        extractor.extract(
            file_path
        )


def test_invalid_docx_file_is_rejected(tmp_path):
    file_path = tmp_path / "invalid.docx"

    file_path.write_text(
        "This is not a valid DOCX package.",
        encoding="utf-8",
    )

    extractor = DocxDocumentExtractor()

    with pytest.raises(
        ValueError,
        match=(
            "not a valid DOCX document"
            "|invalid or corrupted"
            "|Unable to open DOCX document"
        ),
    ):
        extractor.extract(
            file_path
        )


def test_corrupted_docx_package_is_rejected(tmp_path):
    file_path = tmp_path / "corrupted.docx"

    with ZipFile(
        file_path,
        mode="w",
    ) as archive:
        archive.writestr(
            "invalid.xml",
            "<invalid>",
        )

    extractor = DocxDocumentExtractor()

    with pytest.raises(
        ValueError,
        match=(
            "invalid or corrupted"
            "|Unable to open DOCX document"
        ),
    ):
        extractor.extract(
            file_path
        )