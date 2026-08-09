from __future__ import annotations

from pathlib import Path
from typing import Iterable, List, Set
from zipfile import BadZipFile

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from docx.text.paragraph import Paragraph

from backend.app.modules.medical_document_intelligence.contracts.document_content import (
    DocumentContent,
)
from backend.app.modules.medical_document_intelligence.extractors.base import (
    BaseDocumentExtractor,
)


class DocxDocumentExtractor(BaseDocumentExtractor):
    """
    Extract text and structural content from Microsoft Word DOCX files.

    The extractor converts DOCX content into the unified immutable
    DocumentContent contract used by the MedNexus processing pipeline.

    Supported content:
        - Body paragraphs
        - Body tables
        - Section headers
        - Section footers

    Extraction boundaries:
        - No de-identification is performed here.
        - No clinical NLP is performed here.
        - No document classification is performed here.
        - Images and OCR are intentionally deferred to a later stage.
        - Text boxes, comments, tracked changes, and embedded files are
          not yet extracted.

    Body paragraphs and tables are processed in their original document
    order. Headers and footers are appended after the body because they
    belong to section-level document structure rather than the body flow.
    """

    DOCX_MEDIA_TYPE = (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    )

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return (".docx",)

    @property
    def media_type(self) -> str:
        return self.DOCX_MEDIA_TYPE

    def extract(
        self,
        path: str | Path,
    ) -> DocumentContent:
        """
        Extract a DOCX document into a DocumentContent contract.

        Args:
            path:
                Source DOCX file path.

        Returns:
            DocumentContent containing extracted text and metadata.

        Raises:
            FileNotFoundError:
                If the source file does not exist.

            ValueError:
                If the path is not a file, has the wrong extension,
                or does not contain a valid DOCX package.

            PermissionError:
                If the file cannot be read because of operating-system
                permissions.
        """

        file_path = self._validate_path(path)

        try:
            document = Document(file_path)
        except PackageNotFoundError as exc:
            raise ValueError(
                f"Unable to open DOCX document: {file_path.name}."
            ) from exc
        except BadZipFile as exc:
            raise ValueError(
                f"The file is not a valid DOCX document: "
                f"{file_path.name}."
            ) from exc
        except (OSError, KeyError, ValueError) as exc:
            raise ValueError(
                f"The DOCX document is invalid or corrupted: "
                f"{file_path.name}."
            ) from exc

        body_blocks = list(
            self._extract_body_blocks(document)
        )

        header_blocks = self._extract_headers(document)
        footer_blocks = self._extract_footers(document)

        extracted_blocks = [
            *body_blocks,
            *header_blocks,
            *footer_blocks,
        ]

        text = self._join_blocks(extracted_blocks)
        warnings = self._build_warnings(
            text=text,
            document=document,
        )

        metadata = {
            "extractor": "docx",
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "section_count": len(document.sections),
            "header_block_count": len(header_blocks),
            "footer_block_count": len(footer_blocks),
            "body_block_count": len(body_blocks),
            "image_count": len(document.inline_shapes),
            "contains_tables": bool(document.tables),
            "contains_images": bool(document.inline_shapes),
            "preserves_body_order": True,
        }

        return DocumentContent(
            text=text,
            source_name=file_path.name,
            media_type=self.media_type,
            extension=file_path.suffix,
            file_size=file_path.stat().st_size,
            encoding=None,
            page_count=None,
            metadata=metadata,
            warnings=tuple(warnings),
        )

    def _validate_path(
        self,
        path: str | Path,
    ) -> Path:
        """
        Validate and normalize the source DOCX path.
        """

        if not isinstance(path, (str, Path)):
            raise TypeError(
                "path must be a string or pathlib.Path."
            )

        if isinstance(path, str):
            normalized_path = path.strip()

            if not normalized_path:
                raise ValueError(
                    "path must not be empty."
                )

            file_path = Path(normalized_path)
        else:
            file_path = path

        if not file_path.exists():
            raise FileNotFoundError(file_path)

        if not file_path.is_file():
            raise ValueError(
                f"{file_path} is not a file."
            )

        if file_path.suffix.lower() not in self.supported_extensions:
            raise ValueError(
                "DocxDocumentExtractor supports only "
                "'.docx' files."
            )

        return file_path

    def _extract_body_blocks(
        self,
        document: DocxDocument,
    ) -> Iterable[str]:
        """
        Extract body paragraphs and tables in document order.
        """

        for child in document.element.body.iterchildren():

            if isinstance(child, CT_P):
                paragraph = Paragraph(
                    child,
                    document,
                )

                text = self._normalize_text(
                    paragraph.text
                )

                if text:
                    yield text

            elif isinstance(child, CT_Tbl):
                table = Table(
                    child,
                    document,
                )

                table_text = self._extract_table(
                    table
                )

                if table_text:
                    yield table_text

    def _extract_table(
        self,
        table: Table,
    ) -> str:
        """
        Convert a Word table into tab-separated rows.

        Tabs preserve cell boundaries while newlines preserve rows.
        This representation remains readable as text and can later be
        processed by a dedicated table-understanding layer.
        """

        rows: List[str] = []

        for row in table.rows:
            cells = [
                self._normalize_text(cell.text)
                for cell in row.cells
            ]

            if not any(cells):
                continue

            rows.append(
                "\t".join(cells)
            )

        return "\n".join(rows)

    def _extract_headers(
        self,
        document: DocxDocument,
    ) -> List[str]:
        """
        Extract unique section headers.

        Linked headers may be reused by several sections, therefore
        DOCX part names are tracked to prevent duplicated extraction.
        """

        return self._extract_section_parts(
            document=document,
            part_name="header",
        )

    def _extract_footers(
        self,
        document: DocxDocument,
    ) -> List[str]:
        """
        Extract unique section footers.
        """

        return self._extract_section_parts(
            document=document,
            part_name="footer",
        )

    def _extract_section_parts(
        self,
        *,
        document: DocxDocument,
        part_name: str,
    ) -> List[str]:
        """
        Extract unique header or footer content from document sections.
        """

        blocks: List[str] = []
        seen_parts: Set[str] = set()

        for section in document.sections:
            section_part = getattr(
                section,
                part_name,
            )

            part_identifier = str(
                section_part.part.partname
            )

            if part_identifier in seen_parts:
                continue

            seen_parts.add(part_identifier)

            paragraphs = [
                self._normalize_text(paragraph.text)
                for paragraph in section_part.paragraphs
            ]

            paragraphs = [
                paragraph
                for paragraph in paragraphs
                if paragraph
            ]

            table_blocks = [
                self._extract_table(table)
                for table in section_part.tables
            ]

            table_blocks = [
                table_block
                for table_block in table_blocks
                if table_block
            ]

            content_blocks = [
                *paragraphs,
                *table_blocks,
            ]

            if content_blocks:
                blocks.append(
                    "\n".join(content_blocks)
                )

        return blocks

    def _build_warnings(
        self,
        *,
        text: str,
        document: DocxDocument,
    ) -> List[str]:
        """
        Create non-fatal extraction warnings.
        """

        warnings: List[str] = []

        if not text.strip():
            warnings.append(
                "The DOCX document contains no extractable text."
            )

        if document.inline_shapes:
            warnings.append(
                "The DOCX document contains images. "
                "Image text was not extracted because OCR is not "
                "enabled in this stage."
            )

        return warnings

    @staticmethod
    def _normalize_text(
        text: str,
    ) -> str:
        """
        Normalize extracted text while preserving meaningful line breaks.
        """

        if not isinstance(text, str):
            return ""

        lines = [
            line.strip()
            for line in text.replace(
                "\r\n",
                "\n",
            ).replace(
                "\r",
                "\n",
            ).split("\n")
        ]

        return "\n".join(
            line
            for line in lines
            if line
        ).strip()

    @staticmethod
    def _join_blocks(
        blocks: Iterable[str],
    ) -> str:
        """
        Join extracted document blocks using clear paragraph separation.
        """

        normalized_blocks = [
            block.strip()
            for block in blocks
            if isinstance(block, str)
            and block.strip()
        ]

        return "\n\n".join(
            normalized_blocks
        )